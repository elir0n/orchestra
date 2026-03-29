from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_bold_run(paragraph, text: str) -> None:
    """Add a run with bold formatting."""
    run = paragraph.add_run(text)
    run.bold = True


def _parse_inline(paragraph, text: str) -> None:
    """
    Parse inline markdown (**bold**) and add runs to the paragraph.
    Other inline markdown (italics, code) is rendered as plain text.
    """
    # Split on **...**
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 1:
            # Odd indices are bold content (inside **)
            _add_bold_run(paragraph, part)
        else:
            paragraph.add_run(part)


def markdown_to_docx(markdown_text: str, output_path: str) -> None:
    """
    Convert a markdown string (as produced by Claude for CVs) to a .docx file.

    Supported markdown:
    - `# Text`   → Title style
    - `## Text`  → Heading 1
    - `### Text` → Heading 2
    - `- Text`   → List Bullet (with **bold** inline support)
    - blank line → paragraph break (skipped)
    - other line → Normal paragraph (with **bold** inline support)
    """
    doc = Document()

    # Remove default empty paragraph that Word adds
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        if not line:
            continue

        if line.startswith("### "):
            content = line[4:].strip()
            para = doc.add_paragraph(style="Heading 2")
            _parse_inline(para, content)

        elif line.startswith("## "):
            content = line[3:].strip()
            para = doc.add_paragraph(style="Heading 1")
            _parse_inline(para, content)

        elif line.startswith("# "):
            content = line[2:].strip()
            para = doc.add_paragraph(style="Title")
            _parse_inline(para, content)

        elif line.startswith("- ") or line.startswith("* "):
            content = line[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            _parse_inline(para, content)

        else:
            # Plain paragraph — strip any leading/trailing markdown artifacts
            content = line.strip().lstrip(">").strip()
            if content:
                para = doc.add_paragraph(style="Normal")
                _parse_inline(para, content)

    doc.save(output_path)


def read_docx_text(path: str) -> str:
    """Extract plain text from a .docx file (used to read format/master CVs)."""
    doc = Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    return "\n".join(paragraphs)
